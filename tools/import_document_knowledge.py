from __future__ import annotations

import argparse
import re
from pathlib import Path
import sys


ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from sqlalchemy import delete

from mvp_agent.db import get_session, init_db
from mvp_agent.models import KnowledgeChunk
from mvp_agent.vector_store import VectorStoreUnavailable, build_chroma_index


SUPPORTED_SUFFIXES = {".pdf", ".docx", ".txt", ".md"}


def extract_document_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix in {".txt", ".md"}:
        return normalize_text(path.read_text(encoding="utf-8"))
    raise SystemExit(f"Unsupported document type: {path.suffix}. Supported: PDF, DOCX, TXT, MD")


def extract_pdf_text(pdf_path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise SystemExit("Missing pypdf. Install it with: pip install pypdf") from exc

    reader = PdfReader(str(pdf_path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = normalize_text(text)
        if text:
            pages.append(f"第 {index} 页\n{text}")
    return "\n\n".join(pages).strip()


def extract_docx_text(docx_path: Path) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise SystemExit("Missing python-docx. Install it with: pip install python-docx") from exc

    document = Document(str(docx_path))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            blocks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n\n".join(blocks).strip()


def normalize_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_text(text: str, chunk_size: int = 700, overlap: int = 100) -> list[str]:
    paragraphs = [item.strip() for item in re.split(r"\n\s*\n", text) if item.strip()]
    chunks: list[str] = []
    current = ""

    for paragraph in paragraphs:
        if len(paragraph) > chunk_size:
            if current:
                chunks.append(current.strip())
                current = ""
            chunks.extend(split_long_paragraph(paragraph, chunk_size, overlap))
            continue
        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            chunks.append(current.strip())
            prefix = current[-overlap:] if overlap > 0 else ""
            current = f"{prefix}\n\n{paragraph}".strip() if prefix else paragraph

    if current:
        chunks.append(current.strip())
    return chunks


def split_long_paragraph(paragraph: str, chunk_size: int, overlap: int) -> list[str]:
    chunks: list[str] = []
    step = max(1, chunk_size - overlap)
    for start in range(0, len(paragraph), step):
        chunk = paragraph[start : start + chunk_size].strip()
        if chunk:
            chunks.append(chunk)
    return chunks


def import_document(
    path: Path,
    title: str | None = None,
    source: str | None = None,
    chunk_size: int = 700,
    overlap: int = 100,
    replace: bool = True,
    rebuild_index: bool = True,
) -> int:
    init_db()
    path = path.resolve()
    if not path.exists():
        raise SystemExit(f"Document not found: {path}")
    if path.suffix.lower() not in SUPPORTED_SUFFIXES:
        raise SystemExit(f"Unsupported document type: {path.suffix}. Supported: PDF, DOCX, TXT, MD")

    doc_title = title or path.stem
    doc_source = source or path.name
    text = extract_document_text(path)
    if not text:
        raise SystemExit("No text extracted from document. Scanned PDFs need OCR first.")
    chunks = split_text(text, chunk_size=chunk_size, overlap=overlap)

    with get_session() as session:
        if replace:
            session.execute(delete(KnowledgeChunk).where(KnowledgeChunk.source == doc_source))
        for index, chunk in enumerate(chunks, start=1):
            session.add(
                KnowledgeChunk(
                    title=f"{doc_title}-{index:03d}",
                    content=chunk,
                    source=doc_source,
                )
            )

    if rebuild_index:
        try:
            count = build_chroma_index()
            print(f"Rebuilt Chroma index with {count} chunks.")
        except VectorStoreUnavailable as exc:
            print(f"Chroma index was not rebuilt: {exc}")
            print("Install vector dependencies with: pip install chromadb sentence-transformers")
    return len(chunks)


def main() -> None:
    parser = argparse.ArgumentParser(description="Import PDF/DOCX/TXT/MD into RAG knowledge chunks.")
    parser.add_argument("document", help="Path to the document file.")
    parser.add_argument("--title", help="Document title prefix. Defaults to filename.")
    parser.add_argument("--source", help="Source identifier. Defaults to filename.")
    parser.add_argument("--chunk-size", type=int, default=700, help="Approximate max characters per chunk.")
    parser.add_argument("--overlap", type=int, default=100, help="Overlapping characters between chunks.")
    parser.add_argument("--append", action="store_true", help="Append chunks instead of replacing same-source chunks.")
    parser.add_argument("--no-rebuild", action="store_true", help="Skip rebuilding Chroma index.")
    args = parser.parse_args()

    count = import_document(
        path=Path(args.document),
        title=args.title,
        source=args.source,
        chunk_size=args.chunk_size,
        overlap=args.overlap,
        replace=not args.append,
        rebuild_index=not args.no_rebuild,
    )
    print(f"Imported {count} chunks from {args.document}")


if __name__ == "__main__":
    main()
